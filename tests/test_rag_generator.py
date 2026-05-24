"""RAG技术标生成模块测试

测试文档摄入分块、向量检索排序、LLM调用重试、
通用模板后备和Word文档格式化输出。
"""

import os
import tempfile

import pytest

from src.exceptions import UnsupportedFormatError
from src.rag_generator.ingestion import DocumentChunk, DocumentIngester


class TestDocumentIngesterUnsupportedFormat:
    """测试不支持的格式抛出异常"""

    def test_txt_raises_unsupported(self, tmp_path):
        ingester = DocumentIngester()
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("hello")
        with pytest.raises(UnsupportedFormatError):
            ingester.ingest(str(txt_file))

    def test_csv_raises_unsupported(self, tmp_path):
        ingester = DocumentIngester()
        csv_file = tmp_path / "data.csv"
        csv_file.write_text("a,b,c")
        with pytest.raises(UnsupportedFormatError):
            ingester.ingest(str(csv_file))

    def test_no_extension_raises_unsupported(self, tmp_path):
        ingester = DocumentIngester()
        no_ext = tmp_path / "noext"
        no_ext.write_text("content")
        with pytest.raises(UnsupportedFormatError):
            ingester.ingest(str(no_ext))


class TestDocumentIngesterDocx:
    """测试.docx摄入"""

    def _create_docx(self, path, paragraphs, heading=None):
        """Helper to create a .docx file with given paragraphs"""
        from docx import Document

        doc = Document()
        if heading:
            doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)
        doc.save(str(path))

    def test_docx_ingestion_produces_chunks(self, tmp_path):
        docx_path = tmp_path / "sample.docx"
        self._create_docx(docx_path, ["This is paragraph one.", "This is paragraph two."])

        ingester = DocumentIngester()
        chunks = ingester.ingest(str(docx_path))

        assert len(chunks) >= 1
        assert all(isinstance(c, DocumentChunk) for c in chunks)
        assert "paragraph one" in chunks[0].text

    def test_docx_metadata_contains_source_doc(self, tmp_path):
        docx_path = tmp_path / "report.docx"
        self._create_docx(docx_path, ["Some content here."])

        ingester = DocumentIngester()
        chunks = ingester.ingest(str(docx_path))

        assert chunks[0].metadata["source_doc"] == "report.docx"

    def test_docx_heading_extracted_as_section_title(self, tmp_path):
        docx_path = tmp_path / "headed.docx"
        self._create_docx(docx_path, ["Body text."], heading="My Title")

        ingester = DocumentIngester()
        chunks = ingester.ingest(str(docx_path))

        assert chunks[0].metadata["section_title"] == "My Title"

    def test_empty_docx_produces_no_chunks(self, tmp_path):
        docx_path = tmp_path / "empty.docx"
        self._create_docx(docx_path, [])

        ingester = DocumentIngester()
        chunks = ingester.ingest(str(docx_path))

        assert chunks == []


class TestChunkSizeAndOverlap:
    """测试分块大小和重叠"""

    def test_chunk_size_honored(self, tmp_path):
        """Chunks should not exceed chunk_size words"""
        from docx import Document

        docx_path = tmp_path / "long.docx"
        doc = Document()
        # Create a document with 100 words
        words = ["word"] * 100
        doc.add_paragraph(" ".join(words))
        doc.save(str(docx_path))

        ingester = DocumentIngester(chunk_size=30, chunk_overlap=5)
        chunks = ingester.ingest(str(docx_path))

        # Each chunk should have at most 30 words
        for chunk in chunks:
            word_count = len(chunk.text.split())
            assert word_count <= 30

    def test_overlap_creates_shared_content(self, tmp_path):
        """Overlap should cause adjacent chunks to share words"""
        from docx import Document

        docx_path = tmp_path / "overlap.docx"
        doc = Document()
        # Create numbered words so we can verify overlap
        words = [f"w{i}" for i in range(50)]
        doc.add_paragraph(" ".join(words))
        doc.save(str(docx_path))

        ingester = DocumentIngester(chunk_size=20, chunk_overlap=5)
        chunks = ingester.ingest(str(docx_path))

        assert len(chunks) >= 2
        # Last 5 words of chunk 0 should appear at start of chunk 1
        chunk0_words = chunks[0].text.split()
        chunk1_words = chunks[1].text.split()
        overlap_from_first = chunk0_words[-5:]
        overlap_in_second = chunk1_words[:5]
        assert overlap_from_first == overlap_in_second

    def test_multiple_chunks_for_large_doc(self, tmp_path):
        """A document larger than chunk_size should produce multiple chunks"""
        from docx import Document

        docx_path = tmp_path / "big.docx"
        doc = Document()
        words = ["test"] * 200
        doc.add_paragraph(" ".join(words))
        doc.save(str(docx_path))

        ingester = DocumentIngester(chunk_size=50, chunk_overlap=10)
        chunks = ingester.ingest(str(docx_path))

        assert len(chunks) > 1


# ============================================================
# VectorRetriever Tests
# ============================================================
from unittest.mock import patch, MagicMock

from src.exceptions import VectorStoreError
from src.rag_generator.retrieval import RetrievalResult, VectorRetriever


class TestVectorRetrieverStoreAndRetrieve:
    """测试存储和检索返回结果"""

    def test_store_and_retrieve_returns_results(self):
        retriever = VectorRetriever(collection_name="test_store_retrieve")
        chunks = [
            DocumentChunk(text="招标文件要求提供资质证明", metadata={"source_doc": "bid1.docx"}),
            DocumentChunk(text="投标人需具备ISO认证", metadata={"source_doc": "bid2.docx"}),
            DocumentChunk(text="项目预算不超过500万元", metadata={"source_doc": "bid3.docx"}),
        ]
        retriever.store(chunks)

        results = retriever.retrieve("资质要求", top_k=5, threshold=0.0)
        assert len(results) > 0
        assert all(isinstance(r, RetrievalResult) for r in results)
        assert all(isinstance(r.chunk, DocumentChunk) for r in results)

    def test_retrieve_results_have_similarity_scores(self):
        retriever = VectorRetriever(collection_name="test_scores")
        chunks = [
            DocumentChunk(text="Technical requirements for building construction", metadata={"source_doc": "a.docx"}),
        ]
        retriever.store(chunks)

        results = retriever.retrieve("building construction", top_k=5, threshold=0.0)
        assert len(results) > 0
        for r in results:
            assert 0.0 <= r.similarity_score <= 1.0


class TestVectorRetrieverThresholdFiltering:
    """测试阈值过滤移除低相似度结果"""

    def test_high_threshold_filters_low_similarity(self):
        retriever = VectorRetriever(collection_name="test_threshold")
        chunks = [
            DocumentChunk(text="Python programming language tutorial", metadata={"source_doc": "py.docx"}),
            DocumentChunk(text="Java enterprise application development", metadata={"source_doc": "java.docx"}),
            DocumentChunk(text="Cooking recipes for Italian pasta", metadata={"source_doc": "cook.docx"}),
        ]
        retriever.store(chunks)

        # High threshold should filter out unrelated results
        results_strict = retriever.retrieve("Python programming", top_k=5, threshold=0.9)
        results_loose = retriever.retrieve("Python programming", top_k=5, threshold=0.0)

        # Strict threshold should return fewer or equal results
        assert len(results_strict) <= len(results_loose)

    def test_threshold_zero_returns_all(self):
        retriever = VectorRetriever(collection_name="test_threshold_zero")
        chunks = [
            DocumentChunk(text="Document about cats", metadata={"source_doc": "cats.docx"}),
            DocumentChunk(text="Document about dogs", metadata={"source_doc": "dogs.docx"}),
        ]
        retriever.store(chunks)

        results = retriever.retrieve("animals", top_k=5, threshold=0.0)
        # With threshold=0, all stored docs should be returned
        assert len(results) == 2


class TestVectorRetrieverOrdering:
    """测试结果按相似度降序排列"""

    def test_results_ordered_descending_by_similarity(self):
        retriever = VectorRetriever(collection_name="test_ordering")
        chunks = [
            DocumentChunk(text="Machine learning with neural networks deep learning AI", metadata={"source_doc": "ml.docx"}),
            DocumentChunk(text="Gardening tips for growing roses in spring", metadata={"source_doc": "garden.docx"}),
            DocumentChunk(text="Neural network training optimization techniques", metadata={"source_doc": "nn.docx"}),
        ]
        retriever.store(chunks)

        results = retriever.retrieve("deep learning neural networks", top_k=5, threshold=0.0)
        assert len(results) >= 2

        # Verify descending order
        for i in range(len(results) - 1):
            assert results[i].similarity_score >= results[i + 1].similarity_score


class TestVectorRetrieverEmptyCollection:
    """测试空集合返回空列表"""

    def test_empty_collection_returns_empty_list(self):
        retriever = VectorRetriever(collection_name="test_empty_collection")
        results = retriever.retrieve("any query", top_k=5, threshold=0.0)
        assert results == []


class TestVectorRetrieverErrorHandling:
    """测试ChromaDB不可用时抛出VectorStoreError"""

    def test_vectorstore_error_raised_when_chromadb_unavailable(self):
        # Patch sys.modules so the import inside __init__ fails
        with patch.dict("sys.modules", {"chromadb": None}):
            with pytest.raises(VectorStoreError, match="Failed to initialize ChromaDB"):
                VectorRetriever(collection_name="fail_collection")

    def test_retrieve_raises_vectorstore_error_on_query_failure(self):
        retriever = VectorRetriever(collection_name="test_query_error")
        # Mock the collection.query to raise an exception
        retriever.collection.query = MagicMock(side_effect=RuntimeError("DB connection lost"))

        with pytest.raises(VectorStoreError, match="Knowledge base temporarily inaccessible"):
            retriever.retrieve("test query")

    def test_store_raises_vectorstore_error_on_add_failure(self):
        retriever = VectorRetriever(collection_name="test_store_error")
        retriever.collection.add = MagicMock(side_effect=RuntimeError("Write failed"))

        chunks = [DocumentChunk(text="some text", metadata={"source_doc": "x.docx"})]
        with pytest.raises(VectorStoreError, match="Failed to store chunks"):
            retriever.store(chunks)


# ============================================================
# BidGenerator Tests
# ============================================================
from unittest.mock import patch, MagicMock

import httpx

from src.exceptions import LLMAPIError
from src.rag_generator.generation import BidGenerator, GENERIC_TEMPLATE
from src.rag_generator.ingestion import DocumentChunk
from src.rag_generator.retrieval import RetrievalResult


class TestBidGeneratorGenericTemplate:
    """测试无上下文时使用通用模板"""

    def test_generate_with_empty_context_returns_generic_template(self):
        generator = BidGenerator(api_key="test-key", api_url="http://fake")
        result = generator.generate("建设智慧城市项目", context_chunks=[])

        assert "建设智慧城市项目" in result
        assert "技术标文档" in result
        assert "公司简介" in result
        assert "质量保证措施" in result

    def test_generate_with_empty_list_does_not_call_llm(self):
        generator = BidGenerator(api_key="test-key", api_url="http://fake")
        with patch.object(generator, "_call_llm_with_retry") as mock_llm:
            generator.generate("测试项目", context_chunks=[])
            mock_llm.assert_not_called()


class TestBidGeneratorComposePrompt:
    """测试prompt组装包含简介和上下文"""

    def test_compose_prompt_includes_brief(self):
        generator = BidGenerator(api_key="k", api_url="http://x")
        chunks = [
            RetrievalResult(
                chunk=DocumentChunk(text="历史标书内容A", metadata={"source_doc": "bid1.docx"}),
                similarity_score=0.9,
            )
        ]
        prompt = generator._compose_prompt("智慧城市建设", chunks)

        assert "智慧城市建设" in prompt
        assert "历史标书内容A" in prompt
        assert "bid1.docx" in prompt

    def test_compose_prompt_includes_multiple_chunks(self):
        generator = BidGenerator(api_key="k", api_url="http://x")
        chunks = [
            RetrievalResult(
                chunk=DocumentChunk(text="内容一", metadata={"source_doc": "a.docx"}),
                similarity_score=0.9,
            ),
            RetrievalResult(
                chunk=DocumentChunk(text="内容二", metadata={"source_doc": "b.docx"}),
                similarity_score=0.8,
            ),
        ]
        prompt = generator._compose_prompt("项目简介", chunks)

        assert "内容一" in prompt
        assert "内容二" in prompt
        assert "a.docx" in prompt
        assert "b.docx" in prompt


class TestBidGeneratorRetry:
    """测试LLM调用重试机制"""

    @patch("src.rag_generator.generation.time.sleep")
    @patch.object(BidGenerator, "_call_llm")
    def test_retries_on_failure_then_succeeds(self, mock_call_llm, mock_sleep):
        mock_call_llm.side_effect = [
            httpx.HTTPStatusError("500", request=MagicMock(), response=MagicMock()),
            httpx.HTTPStatusError("500", request=MagicMock(), response=MagicMock()),
            "生成的技术标文本",
        ]

        generator = BidGenerator(api_key="k", api_url="http://x", max_retries=3)
        result = generator._call_llm_with_retry("test prompt")

        assert result == "生成的技术标文本"
        assert mock_call_llm.call_count == 3
        # Should have slept twice (after attempt 1 and 2, not after final success)
        assert mock_sleep.call_count == 2

    @patch("src.rag_generator.generation.time.sleep")
    @patch.object(BidGenerator, "_call_llm")
    def test_raises_llm_api_error_after_all_retries_exhausted(self, mock_call_llm, mock_sleep):
        mock_call_llm.side_effect = httpx.ConnectError("Connection refused")

        generator = BidGenerator(api_key="k", api_url="http://x", max_retries=3)

        with pytest.raises(LLMAPIError, match="All 3 retry attempts failed"):
            generator._call_llm_with_retry("test prompt")

        assert mock_call_llm.call_count == 3

    @patch("src.rag_generator.generation.time.sleep")
    @patch.object(BidGenerator, "_call_llm")
    def test_exponential_backoff_timing(self, mock_call_llm, mock_sleep):
        mock_call_llm.side_effect = RuntimeError("fail")

        generator = BidGenerator(api_key="k", api_url="http://x", max_retries=3)

        with pytest.raises(LLMAPIError):
            generator._call_llm_with_retry("prompt")

        # Backoff: 2^0=1s, 2^1=2s (no sleep after last attempt)
        mock_sleep.assert_any_call(1)
        mock_sleep.assert_any_call(2)
        assert mock_sleep.call_count == 2


# ============================================================
# WordFormatter Tests
# ============================================================
from src.rag_generator.formatter import WordFormatter


class TestWordFormatterOutput:
    """测试WordFormatter生成有效的.docx文件"""

    def test_format_creates_valid_docx_file(self, tmp_path):
        formatter = WordFormatter()
        project_info = {"project_name": "智慧城市", "date": "2024-01-15"}
        generated_text = "# 公司简介\n我司成立于2000年。\n# 技术方案\n采用先进技术。"

        result_path = formatter.format(generated_text, project_info, str(tmp_path))

        assert os.path.isfile(result_path)
        assert result_path.endswith(".docx")

    def test_output_file_path_returned_correctly(self, tmp_path):
        formatter = WordFormatter()
        project_info = {"project_name": "测试项目", "date": "2024-06-01"}

        result_path = formatter.format("一些文本", project_info, str(tmp_path))

        assert "测试项目" in result_path
        assert str(tmp_path) in result_path


class TestWordFormatterSections:
    """测试生成文档包含所有必需章节"""

    def test_document_contains_all_required_sections(self, tmp_path):
        from docx import Document as DocxDocument

        formatter = WordFormatter()
        project_info = {"project_name": "工程项目", "date": "2024-03-01"}
        generated_text = (
            "## 公司简介\n公司成立于2005年。\n"
            "## 技术方案\n采用BIM技术。\n"
            "## 项目进度计划\n工期6个月。\n"
            "## 质量保证措施\nISO9001认证。"
        )

        result_path = formatter.format(generated_text, project_info, str(tmp_path))
        doc = DocxDocument(result_path)

        # Collect all heading text
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

        assert "目录" in headings
        assert "公司简介" in headings
        assert "技术方案" in headings
        assert "项目进度计划" in headings
        assert "质量保证措施" in headings

    def test_sections_without_content_show_placeholder(self, tmp_path):
        from docx import Document as DocxDocument

        formatter = WordFormatter()
        project_info = {"project_name": "空白项目", "date": "2024-01-01"}

        result_path = formatter.format("", project_info, str(tmp_path))
        doc = DocxDocument(result_path)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[待补充]" in all_text


class TestWordFormatterCoverPage:
    """测试封面页包含项目名称"""

    def test_cover_page_includes_project_name(self, tmp_path):
        from docx import Document as DocxDocument

        formatter = WordFormatter()
        project_info = {"project_name": "智慧交通建设", "date": "2024-05-20"}

        result_path = formatter.format("技术内容", project_info, str(tmp_path))
        doc = DocxDocument(result_path)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "智慧交通建设" in all_text

    def test_cover_page_includes_date(self, tmp_path):
        from docx import Document as DocxDocument

        formatter = WordFormatter()
        project_info = {"project_name": "项目A", "date": "2024-12-25"}

        result_path = formatter.format("内容", project_info, str(tmp_path))
        doc = DocxDocument(result_path)

        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "2024-12-25" in all_text
