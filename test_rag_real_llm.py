"""Live test: RAG bid generation with real LLM API"""
import sys, json, logging
sys.path.insert(0, "D:/CADAI")
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

from pathlib import Path
from docx import Document as DocxDoc

# Step 1: Prepare historical bid document
print("="*60)
print("Step 1: Preparing historical bid documents")
print("="*60)

hist_dir = Path("D:/CADAI/data/historical_bids")
hist_dir.mkdir(parents=True, exist_ok=True)

# Create a richer historical bid document
sample_path = hist_dir / "market_construction_bid_2025.docx"
doc = DocxDoc()
doc.add_heading("某市政道路改造工程技术标", level=1)

doc.add_heading("公司简介", level=2)
doc.add_paragraph(
    "我司成立于2010年，注册资金5000万元，具有市政公用工程施工总承包壹级资质、"
    "建筑工程施工总承包壹级资质、装饰装修工程专业承包壹级资质。"
    "公司现有员工260人，其中高级工程师35人、一级建造师22人、二级建造师45人。"
    "近五年承建市政工程项目120余项，累计合同额超过15亿元。"
    "代表性项目包括：XX市南环路快速化改造工程（造价2.3亿元）、"
    "XX新区基础设施一期工程（造价1.8亿元）、XX产业园区配套道路工程（造价9600万元）。"
    "公司连续五年获评AAA级信用企业，获得省级优质工程奖8项。"
)

doc.add_heading("技术方案", level=2)
doc.add_paragraph(
    "一、施工总体部署\n"
    "本项目采用BIM技术进行全过程管理，实行项目经理负责制。"
    "施工划分为三个标段同步推进，关键节点设置里程碑考核。\n\n"
    "二、主要施工方案\n"
    "1. 路基工程：采用强夯法加固地基，夯击能量3000kN·m，处理深度8-10m。"
    "填方路段采用分层碾压工艺，每层松铺厚度不超过30cm，压实度≥96%。\n"
    "2. 路面工程：采用SBS改性沥青混凝土路面，结构层总厚度72cm。"
    "面层为4cm AC-13C + 6cm AC-20C，基层为20cm水泥稳定碎石 + 20cm级配碎石。\n"
    "3. 排水工程：雨水管道采用HDPE双壁波纹管，管径DN300-DN800。"
    "采用开槽施工，槽底设10cm碎石垫层，管道安装后立即进行闭水试验。\n"
    "4. 桥梁工程：采用预应力混凝土连续箱梁结构，跨径组合3×30m。"
    "支架现浇施工，支架地基承载力不低于150kPa。\n\n"
    "三、质量控制要点\n"
    "- 原材料100%送检合格后方可使用\n"
    "- 隐蔽工程验收合格后方可进入下道工序\n"
    "- 关键工序实行三检制（自检、互检、专检）\n"
    "- 成品保护措施到位，设置专人巡查"
)

doc.add_heading("项目进度计划", level=2)
doc.add_paragraph(
    "总工期240天（8个月），分四个阶段：\n"
    "第一阶段（1-2月）：场地准备、管线迁改、临建搭设，投入人员50人\n"
    "第二阶段（3-5月）：路基施工、排水管道安装、桥梁基础，投入人员120人\n"
    "第三阶段（5-7月）：路面结构层施工、桥梁上部结构，投入人员150人\n"
    "第四阶段（7-8月）：附属工程、绿化恢复、竣工验收，投入人员80人\n\n"
    "关键路径：路基施工→路面基层→路面面层→交工验收\n"
    "采用P6软件编制网络计划，每周进行进度偏差分析并及时纠偏。"
)

doc.add_heading("质量保证措施", level=2)
doc.add_paragraph(
    "1. 质量管理体系：执行ISO9001:2015质量管理体系标准，"
    "设置项目质量总监、质量工程师各1名，专职质检员4名。\n"
    "2. 过程控制：实行PDCA循环管理，每道工序设置质量检查点。"
    "混凝土施工前进行配合比试验，沥青路面施工前进行试验段铺筑。\n"
    "3. 检测措施：购置先进检测设备，路面平整度采用3m直尺逐板检查，"
    "压实度采用灌砂法现场检测，弯沉值采用贝克曼梁检测。\n"
    "4. 成品保护：新铺路面设置围挡和警示标志，养生期不少于14天。"
    "桥梁混凝土浇筑后采用土工布覆盖洒水养护，养护时间不少于7天。\n"
    "5. 质量目标：工程合格率100%，优良率≥90%，争创省级优质工程。"
)

doc.save(str(sample_path))
print(f"Created: {sample_path}")

# Step 2: Ingest into ChromaDB
print(f"\n{'='*60}")
print("Step 2: Ingesting into ChromaDB")
print("="*60)

from src.rag_generator.ingestion import DocumentIngester
from src.rag_generator.retrieval import VectorRetriever

ingester = DocumentIngester(chunk_size=200, chunk_overlap=20)
chunks = ingester.ingest(str(sample_path))
print(f"  Chunks created: {len(chunks)}")

retriever = VectorRetriever(collection_name="bid_docs_real_test")
retriever.store(chunks)
print(f"  Stored in ChromaDB")

# Step 3: Retrieve relevant context
print(f"\n{'='*60}")
print("Step 3: Retrieving relevant context for new project")
print("="*60)

project_brief = (
    "海南海事局青年公寓修缮项目，预算395.14万元，工期120天。"
    "项目内容：办公用房修缮施工，包含建筑装饰装修、给排水改造、"
    "电气线路更新、消防系统升级。要求具备建筑工程施工总承包乙级及以上资质。"
)

results = retriever.retrieve(project_brief, top_k=5, threshold=0.0)
print(f"  Retrieved {len(results)} chunks")
for i, r in enumerate(results):
    print(f"  [{i+1}] score={r.similarity_score:.3f}: {r.chunk.text[:60]}...")

# Step 4: Generate with real LLM
print(f"\n{'='*60}")
print("Step 4: Generating with real LLM API (jucode.cn)")
print("="*60)

from src.rag_generator.generation import BidGenerator

generator = BidGenerator(
    api_key="sk-juc-sIcc5XvwTe9zEm3EMYK8BnUM5noiHiStehQ87t1M178",
    api_url="https://api.jucode.cn/v1",
    model="gpt-5.4",
    timeout=120,
    max_retries=2,
)

print("  Calling LLM API...")
generated_text = generator.generate(project_brief, context_chunks=results)
print(f"  Generated {len(generated_text)} characters")
print(f"\n  --- Preview (first 1000 chars) ---")
print(generated_text[:1000])

# Step 5: Format into Word
print(f"\n{'='*60}")
print("Step 5: Formatting into Word document")
print("="*60)

from src.rag_generator.formatter import WordFormatter

formatter = WordFormatter()
output_path = formatter.format(
    generated_text,
    {"project_name": "海南海事局青年公寓修缮项目", "date": "2026-05-24"},
    "D:/CADAI/output"
)
print(f"  Output: {output_path}")

# Verify
from docx import Document as VerifyDoc
vdoc = VerifyDoc(output_path)
total_text = "\n".join(p.text for p in vdoc.paragraphs if p.text.strip())
print(f"  Total content: {len(total_text)} characters")
headings = [p.text for p in vdoc.paragraphs if p.style.name.startswith("Heading")]
print(f"  Sections: {headings}")

print(f"\n{'='*60}")
print("DONE! Full RAG pipeline with real LLM completed.")
print(f"{'='*60}")
