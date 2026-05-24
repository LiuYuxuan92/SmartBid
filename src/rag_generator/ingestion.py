"""文档摄入与分块模块"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.exceptions import UnsupportedFormatError

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {".docx", ".pdf"}


@dataclass
class DocumentChunk:
    """文档分块"""
    text: str
    metadata: dict  # {source_doc, section_title, doc_date}
    embedding: Optional[list[float]] = None


class DocumentIngester:
    """文档摄入器 - 支持.docx和.pdf格式"""

    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 50):
        self.chunk_size = chunk_size  # tokens (approximate by words)
        self.chunk_overlap = chunk_overlap

    def ingest(self, file_path: str) -> list[DocumentChunk]:
        """摄入文档，分块返回

        Args:
            file_path: 文档路径

        Returns:
            DocumentChunk列表

        Raises:
            UnsupportedFormatError: 非.docx/.pdf格式
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in SUPPORTED_FORMATS:
            raise UnsupportedFormatError(
                f"Unsupported format '{suffix}'. Supported: .docx, .pdf"
            )

        if suffix == ".docx":
            text, metadata = self._extract_text_docx(file_path)
        else:
            text, metadata = self._extract_text_pdf(file_path)

        chunks = self._split_into_chunks(text, metadata)
        return chunks

    def _extract_text_docx(self, path: str) -> tuple[str, dict]:
        """从docx提取文本和元数据"""
        from docx import Document

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        # Extract metadata
        metadata = {
            "source_doc": Path(path).name,
            "section_title": "",
            "doc_date": "",
        }

        # Try to get title from first heading
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading"):
                metadata["section_title"] = p.text
                break

        return text, metadata

    def _extract_text_pdf(self, path: str) -> tuple[str, dict]:
        """从pdf提取文本和元数据 (MVP简化实现)"""
        try:
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = [page.extract_text() or "" for page in reader.pages]
                text = "\n".join(pages)
        except ImportError:
            # Fallback if PyPDF2 not available
            logger.warning("PyPDF2 not installed, cannot extract PDF text")
            text = ""

        metadata = {
            "source_doc": Path(path).name,
            "section_title": "",
            "doc_date": "",
        }
        return text, metadata

    def _split_into_chunks(self, text: str, metadata: dict) -> list[DocumentChunk]:
        """按token数分块，保持overlap

        使用空格分词近似token计数 (中文按字符)
        """
        # Simple tokenization: split by whitespace, treating Chinese chars as individual tokens
        words = text.split()

        if not words:
            return []

        chunks = []
        start = 0

        while start < len(words):
            end = min(start + self.chunk_size, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            chunk = DocumentChunk(
                text=chunk_text,
                metadata=metadata.copy(),
            )
            chunks.append(chunk)

            # Advance with overlap
            start = end - self.chunk_overlap
            if start >= len(words):
                break
            # Prevent infinite loop if overlap >= chunk_size
            if end == len(words):
                break

        return chunks
