"""Word文档格式化模块

负责将 LLM 生成的技术标文本格式化为 .docx 文档，
包含封面、目录、公司简介、技术方案、项目进度和质量保证章节。
"""

import logging
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class WordFormatter:
    """Word文档格式化器 - 将生成文本格式化为标准技术标文档"""

    SECTIONS = [
        "cover_page",
        "table_of_contents",
        "company_introduction",
        "technical_approach",
        "project_timeline",
        "quality_assurance",
    ]

    def format(self, generated_text: str, project_info: dict, output_dir: str) -> str:
        """格式化为Word文档

        Args:
            generated_text: LLM生成的技术标文本
            project_info: {"project_name": str, "date": str, ...}
            output_dir: 输出目录

        Returns:
            生成文档的文件路径
        """
        doc = Document()

        self._create_cover_page(doc, project_info)
        self._add_toc(doc)
        self._add_sections(doc, generated_text)

        # Save
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        filename = f"bid_document_{project_info.get('project_name', 'draft')}_{date.today().isoformat()}.docx"
        file_path = output_path / filename
        doc.save(str(file_path))

        logger.info(f"Word document generated: {file_path}")
        return str(file_path)

    def _create_cover_page(self, doc, project_info: dict) -> None:
        """创建封面页"""
        # Add blank lines for spacing
        for _ in range(4):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(project_info.get("project_name", "技术标文档"))
        run.font.size = Pt(28)
        run.bold = True

        doc.add_paragraph()

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("技术标")
        run.font.size = Pt(20)

        doc.add_paragraph()
        doc.add_paragraph()

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(project_info.get("date", date.today().isoformat()))

        doc.add_page_break()

    def _add_toc(self, doc) -> None:
        """添加目录占位"""
        doc.add_heading("目录", level=1)
        doc.add_paragraph("（目录将在文档最终排版时自动生成）")
        doc.add_page_break()

    def _add_sections(self, doc, text: str) -> None:
        """添加正文各章节"""
        section_titles = {
            "company_introduction": "公司简介",
            "technical_approach": "技术方案",
            "project_timeline": "项目进度计划",
            "quality_assurance": "质量保证措施",
        }

        # Try to split generated text by section headers
        sections = self._split_text_into_sections(text)

        for key, title in section_titles.items():
            doc.add_heading(title, level=1)
            content = sections.get(key, "")
            if content:
                for para in content.split("\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph("[待补充]")

    def _split_text_into_sections(self, text: str) -> dict:
        """尝试将生成文本按章节拆分"""
        sections = {}
        current_key = None
        current_content = []

        section_markers = {
            "公司简介": "company_introduction",
            "技术方案": "technical_approach",
            "项目进度计划": "project_timeline",
            "质量保证措施": "quality_assurance",
        }

        for line in text.split("\n"):
            matched = False
            for marker, key in section_markers.items():
                if marker in line and (line.startswith("#") or line.startswith("##")):
                    if current_key and current_content:
                        sections[current_key] = "\n".join(current_content)
                    current_key = key
                    current_content = []
                    matched = True
                    break

            if not matched and current_key:
                current_content.append(line)

        if current_key and current_content:
            sections[current_key] = "\n".join(current_content)

        # If no sections were parsed, put all text in technical_approach
        if not sections and text.strip():
            sections["technical_approach"] = text

        return sections
